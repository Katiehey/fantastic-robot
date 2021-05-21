from docx import Document
from docx.shared import Inches
import pyttsx3

# pyttsx3.speak('hello')


def speak(text):
    pyttsx3.speak(text)


document = Document()

# picture
document.add_picture(
    'I AM WOMAN.jpg',
    width=Inches(2.0))

# name, phone number and email details
name = input('What is your name? ')
speak('hello ' + name + ' how are you today? ')

speak('What is your phone number? ')
phone_number = input('What is your phone number? ')
email = 'KUTLWANOMELAMU93@GMAIL.COM'

document.add_paragraph(
    name + ' : ' +
    phone_number + ' : ' +
    email)

# about me
document.add_heading('About me')
# CAN DO IT THIS WAY OR...
# about_me = input('Tell me about yourself ')
# document.add_paragraph(about_me)

# THIS WAY
document.add_paragraph(input('Tell me about yourself' + ' '))

# WORK EXPERIENCE
document.add_heading('Work experience')
p = document.add_paragraph()

company = input('Enter company name ')
from_date = input('From date ')
to_date = input('To date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('How was your experience at ' + company + ' ')
p.add_run(experience_details)

# MORE EXPERIENCES
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company name ')
        from_date = input('From date ')
        to_date = input('To date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'How was your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# SKILLS
document.add_heading('Skills')
skill = input('Enter skill' + ' ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No' + ' ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill' + ' ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'

    else:
        break

# FOOTER
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Amigoscode and Intuit QuickBooks course project'


document.save('cv.docx')
